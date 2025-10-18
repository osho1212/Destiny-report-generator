from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Drawing, Line
from PIL import Image as PILImage
from openpyxl import Workbook
from jinja2 import Template
from datetime import datetime
import os
import html
import base64
import io
from PyPDF2 import PdfReader, PdfWriter

class ReportGenerator:
    def __init__(self):
        self.output_dir = 'generated_reports'
        os.makedirs(self.output_dir, exist_ok=True)

    def sanitize_input(self, text):
        """Sanitize user input to prevent XSS and injection attacks"""
        if not text:
            return ''
        return html.escape(str(text))

    def base64_to_image(self, base64_string, max_width=6*inch):
        """Convert base64 string to reportlab Image object"""
        try:
            # Remove data URL prefix if present
            if ',' in base64_string:
                base64_string = base64_string.split(',')[1]

            # Decode base64
            image_data = base64.b64decode(base64_string)

            # Open with PIL to get dimensions and convert if needed
            pil_image = PILImage.open(io.BytesIO(image_data))

            # Convert RGBA to RGB if necessary
            if pil_image.mode in ('RGBA', 'P'):
                rgb_image = PILImage.new('RGB', pil_image.size, (255, 255, 255))
                if pil_image.mode == 'P':
                    pil_image = pil_image.convert('RGBA')
                if pil_image.mode == 'RGBA':
                    rgb_image.paste(pil_image, mask=pil_image.split()[3])
                else:
                    rgb_image.paste(pil_image)
                pil_image = rgb_image

            # Save to BytesIO
            img_io = io.BytesIO()
            pil_image.save(img_io, format='JPEG', quality=85)
            img_io.seek(0)

            # Create reportlab Image
            img = Image(img_io)

            # Calculate aspect ratio and resize if needed
            aspect = pil_image.height / pil_image.width
            if img.drawWidth > max_width:
                img.drawWidth = max_width
                img.drawHeight = max_width * aspect

            return img
        except Exception as e:
            print(f"Error converting base64 to image: {str(e)}")
            return None

    def format_dasha(self, form_data, prefix):
        """Format dasha data from form fields
        Format: Planet(Source), NL(NL Source), (Additional house), SL(SL source)
        If no_star is True for the dasha, add * as superscript to the first planet
        """
        planet = self.sanitize_input(form_data.get(f'{prefix}_planet', ''))
        source = self.sanitize_input(form_data.get(f'{prefix}_source', ''))
        nl = self.sanitize_input(form_data.get(f'{prefix}_nl', ''))
        nl_source = self.sanitize_input(form_data.get(f'{prefix}_nl_source', ''))
        additional_house = self.sanitize_input(form_data.get(f'{prefix}_additional_house', ''))
        sl = self.sanitize_input(form_data.get(f'{prefix}_sl', ''))
        sl_source = self.sanitize_input(form_data.get(f'{prefix}_sl_source', ''))
        no_star = form_data.get(f'{prefix}_no_star', False)

        parts = []

        # Planet(Source) - add * if no_star is True
        if planet:
            planet_text = planet
            if no_star:
                planet_text = f'{planet}*'

            if source:
                parts.append(f'{planet_text}({source})')
            else:
                parts.append(planet_text)

        # NL Planet(NL Source)
        if nl:
            if nl_source:
                parts.append(f'{nl}({nl_source})')
            else:
                parts.append(nl)

        # (Additional house)
        if additional_house:
            parts.append(f'({additional_house})')

        # SL Planet(SL source)
        if sl:
            if sl_source:
                parts.append(f'{sl}({sl_source})')
            else:
                parts.append(sl)

        # Add period dates if this is pratyantardasha
        if prefix == 'pratyantardasha':
            period_from = self.sanitize_input(form_data.get(f'{prefix}_period_from', ''))
            period_to = self.sanitize_input(form_data.get(f'{prefix}_period_to', ''))

            if period_from and period_to:
                # Format dates
                from datetime import datetime as dt
                try:
                    from_date = dt.strptime(period_from, '%Y-%m-%d').strftime('%b %d, %Y')
                    to_date = dt.strptime(period_to, '%Y-%m-%d').strftime('%b %d, %Y')
                    parts.append(f'Period: {from_date} - {to_date}')
                except:
                    # If date parsing fails, use raw values
                    parts.append(f'Period: {period_from} - {period_to}')

        return ', '.join(parts) if parts else ''

    def generate_pdf(self, form_data):
        """Generate PDF Destiny Report from form data"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name = self.sanitize_input(form_data.get('name', 'Client')).replace(' ', '_')
        filename = f'destiny_report_{client_name}_{timestamp}.pdf'
        filepath = os.path.join(self.output_dir, filename)

        # Create PDF
        doc = SimpleDocTemplate(filepath, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # Main Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=1,
            fontName='Times-Bold'
        )
        story.append(Paragraph("DESTINY REPORT", title_style))
        story.append(Spacer(1, 0.5*inch))

        # Section Title Style - orange color matching preview, Times New Roman
        section_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#ff8c00'),
            spaceBefore=30,
            spaceAfter=16,
            fontName='Times-Bold'
        )

        # Field Label Style - Times New Roman, bold
        field_label_style = ParagraphStyle(
            'FieldLabel',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=4,
            fontName='Times-Bold',
            textColor=colors.HexColor('#333333')
        )

        # Field Value Style - Times New Roman, with bullet points
        field_value_style = ParagraphStyle(
            'FieldValue',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            fontName='Times-Roman',
            leftIndent=20,
            bulletIndent=10,
            bulletFontName='Times-Roman'
        )

        # Field Style - Times New Roman (for combined label+value)
        field_style = ParagraphStyle(
            'FieldStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            fontName='Times-Roman',
            leftIndent=15
        )

        # Helper function to format multi-line values with bullet points
        def format_value_with_bullets(value):
            """Format values that contain multiple lines or comma-separated items with bullets"""
            if not value:
                return []

            # Check if value contains newlines or commas suggesting multiple items
            if '\n' in value:
                items = [item.strip() for item in value.split('\n') if item.strip()]
            elif ',' in value and len(value) > 50:  # Only split by comma if text is long
                items = [item.strip() for item in value.split(',') if item.strip()]
            else:
                # Single value, no bullets needed
                return [value]

            return items if len(items) > 1 else [value]

        # Helper function to check if text already has bullet points
        def has_bullet(text):
            """Check if text starts with a bullet point"""
            if not text:
                return False
            text = text.strip()
            # Check for common bullet characters
            bullet_chars = ['•', '◦', '○', '●', '▪', '▫', '–', '-', '*', '→']
            return any(text.startswith(char) for char in bullet_chars) or (len(text) > 2 and text[0].isdigit() and text[1] in ['.', ')'])

        # Helper function to add section (only if it has content)
        def add_section(title, fields, keep_inline=False):
            # Check if any field has a value
            has_content = any(self.sanitize_input(form_data.get(field_key, '')) for field_key, _ in fields)

            if not has_content:
                return  # Skip empty sections

            story.append(Paragraph(title, section_style))

            # Add orange line under section title
            d = Drawing(6.5*inch, 2)
            line = Line(0, 1, 6.5*inch, 1)
            line.strokeColor = colors.HexColor('#ff8c00')
            line.strokeWidth = 2
            d.add(line)
            story.append(d)
            story.append(Spacer(1, 0.15*inch))

            for field_key, field_label in fields:
                value = self.sanitize_input(form_data.get(field_key, ''))
                if value:
                    if keep_inline:
                        # Keep label and value on same line (for About the Client)
                        text = f"<b>{field_label}:</b> {value}"
                        story.append(Paragraph(text, field_style))
                    else:
                        # Add field label
                        story.append(Paragraph(f"<b>{field_label}:</b>", field_label_style))

                        # Format value with bullets if multi-line
                        items = format_value_with_bullets(value)
                        if len(items) > 1:
                            # Multiple items - use bullets only if not already present
                            for item in items:
                                if has_bullet(item):
                                    # Already has bullet, don't add another
                                    story.append(Paragraph(item, field_value_style))
                                else:
                                    # Add bullet
                                    bullet_text = f"• {item}"
                                    story.append(Paragraph(bullet_text, field_value_style))
                        else:
                            # Single item - no bullet
                            story.append(Paragraph(items[0], field_value_style))

                    story.append(Spacer(1, 0.08*inch))

            story.append(Spacer(1, 0.1*inch))

        # ABOUT THE CLIENT
        add_section("ABOUT THE CLIENT", [
            ('name', 'Name'),
            ('dateOfBirth', 'Date of Birth'),
            ('timeOfBirth', 'Time of Birth'),
            ('placeOfBirth', 'Place of Birth')
        ])

        # VASTU ANALYSIS
        add_section("VASTU ANALYSIS", [
            ('mapOfHouse', 'Map of the House'),
            ('vastuAnalysis', 'Analysis (Entrances, Kitchen, Washrooms)'),
            ('vastuRemedies', 'Remedies')
        ])


        # Add house map images and room directions if present
        house_map_images = form_data.get('houseMapImages', [])
        house_map_analyses = form_data.get('houseMapAnalyses', [])

        if house_map_images and isinstance(house_map_images, list):
            for idx, img_base64 in enumerate(house_map_images):
                img = self.base64_to_image(img_base64)
                if img:
                    story.append(img)
                    story.append(Spacer(1, 0.1*inch))

                    # Add room directions for this map if available
                    if idx < len(house_map_analyses) and house_map_analyses[idx]:
                        story.append(Paragraph(f"<b>Room Directions (Map {idx + 1}):</b>", field_label_style))
                        # Split room directions into individual lines and create separate paragraphs
                        room_lines = house_map_analyses[idx].split('\n')
                        for room_line in room_lines:
                            if room_line.strip():
                                story.append(Paragraph(room_line, field_value_style))
                        story.append(Spacer(1, 0.15*inch))
        # ASTROLOGY (Custom handling for Dashas)
        # Check if ASTROLOGY section has any content
        mahadasha = self.format_dasha(form_data, 'mahadasha')
        antardasha = self.format_dasha(form_data, 'antardasha')
        pratyantardasha = self.format_dasha(form_data, 'pratyantardasha')

        astrology_fields = [
            'donationsToDo', 'donationsToWhom', 'gemstones', 'mantra',
            'birthNakshatra', 'mobileDisplayPicture', 'beneficialSymbols',
            'nakshatraProsperitySymbols', 'nakshatraMentalPhysicalWellbeing',
            'nakshatraAccomplishments', 'nakshatraAvoidSymbols'
        ]

        has_astrology_content = (
            mahadasha or antardasha or pratyantardasha or
            any(self.sanitize_input(form_data.get(field, '')) for field in astrology_fields)
        )

        if has_astrology_content:
            story.append(Paragraph("ASTROLOGY", section_style))

            # Add orange line under section title
            d = Drawing(6.5*inch, 2)
            line = Line(0, 1, 6.5*inch, 1)
            line.strokeColor = colors.HexColor('#ff8c00')
            line.strokeWidth = 2
            d.add(line)
            story.append(d)
            story.append(Spacer(1, 0.15*inch))

            # Add Mahadasha
            if mahadasha:
                story.append(Paragraph("<b>Mahadasha:</b>", field_label_style))
                story.append(Paragraph(mahadasha, field_value_style))
                story.append(Spacer(1, 0.08*inch))

            # Add Antardasha
            if antardasha:
                story.append(Paragraph("<b>Antardasha:</b>", field_label_style))
                story.append(Paragraph(antardasha, field_value_style))
                story.append(Spacer(1, 0.08*inch))

            # Add Pratyantardasha
            if pratyantardasha:
                story.append(Paragraph("<b>Pratyantar Dasha:</b>", field_label_style))
                story.append(Paragraph(pratyantardasha, field_value_style))
                story.append(Spacer(1, 0.08*inch))

            # Add other astrology fields
            for field_key, field_label in [
                ('donationsToDo', 'Donations to Do'),
                ('donationsToWhom', 'Donations - To Whom'),
                ('gemstones', 'Gemstones'),
                ('mantra', 'Mantra to Make Situation Positive'),
                ('birthNakshatra', 'Your Birth Nakshatra'),
                ('mobileDisplayPicture', 'Beneficial Mobile Display Picture'),
                ('beneficialSymbols', 'Most Beneficial Symbols'),
                ('nakshatraProsperitySymbols', 'Prosperity Giving Symbols'),
                ('nakshatraMentalPhysicalWellbeing', 'Mental/Physical Wellbeing Symbols'),
                ('nakshatraAccomplishments', 'Accomplishment/Achievement Symbols'),
                ('nakshatraAvoidSymbols', 'Symbols to Avoid')
            ]:
                value = self.sanitize_input(form_data.get(field_key, ''))
                if value:
                    # Add field label
                    story.append(Paragraph(f"<b>{field_label}:</b>", field_label_style))

                    # Format value with bullets if multi-line
                    items = format_value_with_bullets(value)
                    if len(items) > 1:
                        # Multiple items - use bullets only if not already present
                        for item in items:
                            if has_bullet(item):
                                # Already has bullet, don't add another
                                story.append(Paragraph(item, field_value_style))
                            else:
                                # Add bullet
                                bullet_text = f"• {item}"
                                story.append(Paragraph(bullet_text, field_value_style))
                    else:
                        # Single item - no bullet
                        story.append(Paragraph(items[0], field_value_style))

                    story.append(Spacer(1, 0.08*inch))

            story.append(Spacer(1, 0.1*inch))

        # ASTRO VASTU SOLUTION
        add_section("ASTRO VASTU SOLUTION", [
            ('aspectsOnHouses', 'Aspects on Houses'),
            ('aspectsOnPlanets', 'Aspects on Planets'),
            ('whatToRemove', 'What to Remove from Which Directions'),
            ('whatToPlace', 'What to Place in Which Directions'),
            ('astroVastuRemediesBody', 'Astro Vastu Remedies for Body'),
            ('colorObjectsToUse', 'What Color or Objects to Use on Body'),
            ('colorObjectsNotToUse', 'What Color or Objects NOT to Use on Body'),
            ('lockerLocation', 'Most Favorable Location of Locker'),
            ('laughingBuddhaDirection', 'Placement of Laughing Buddha/Vision Board - Direction'),
            ('wishListInkColor', 'Color of Ink to Write Wish List')
        ])

        # GUIDELINES
        add_section("GUIDELINES", [
            ('neverCriticize', 'To Whom You Should Never Criticize or Judge'),
            ('importantBooks', 'Important Books to Read to Uplift Your Life'),
            ('giftsToGive', 'Gifts - To Give'),
            ('giftsToReceive', 'Gifts - To Receive')
        ])

        # BHRIGUNANDA NADI
        add_section("BHRIGUNANDA NADI", [
            ('saturnRelation', 'Saturn Relation'),
            ('saturnFollowing', 'Saturn is following'),
            ('professionalMindset', 'Professional Mindset'),
            ('venusRelation', 'Venus Relation'),
            ('venusFollowing', 'Venus is following'),
            ('financialMindset', 'Financial Mindset')
        ])

        # Add timestamp
        story.append(Spacer(1, 0.3*inch))
        timestamp_text = f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        story.append(Paragraph(timestamp_text, styles['Italic']))

        # Build PDF
        doc.build(story)

        # If Kundli PDF is provided, merge specific pages (1, 3, 4) at the beginning
        kundli_pdf_data = form_data.get('kundliPdf')
        kundli_pages = [1, 3, 4]  # Pages to extract from Kundli PDF

        print(f"Checking for Kundli PDF... Found: {bool(kundli_pdf_data)}")
        if kundli_pdf_data:
            print(f"Kundli PDF data starts with: {kundli_pdf_data[:50] if len(kundli_pdf_data) > 50 else kundli_pdf_data}")

        if kundli_pdf_data and kundli_pdf_data.startswith('data:application/pdf'):
            try:
                print(f"Merging Kundli pages {kundli_pages} at the beginning...")

                # Extract base64 data
                pdf_data = kundli_pdf_data.split(',')[1]
                pdf_bytes = base64.b64decode(pdf_data)

                # Read the Kundli PDF
                kundli_reader = PdfReader(io.BytesIO(pdf_bytes))
                print(f"Kundli PDF has {len(kundli_reader.pages)} pages")

                # Read the generated report PDF
                report_reader = PdfReader(filepath)
                print(f"Generated report has {len(report_reader.pages)} pages")

                # Create a new PDF writer
                pdf_writer = PdfWriter()

                # Add selected Kundli pages first (1-indexed to 0-indexed)
                for page_num in kundli_pages:
                    if 1 <= page_num <= len(kundli_reader.pages):
                        pdf_writer.add_page(kundli_reader.pages[page_num - 1])
                        print(f"Added Kundli page {page_num}")

                # Add all pages from the generated report
                for page in report_reader.pages:
                    pdf_writer.add_page(page)

                # Write the merged PDF
                with open(filepath, 'wb') as output_file:
                    pdf_writer.write(output_file)

                print(f"Successfully merged Kundli pages {kundli_pages} at the beginning")
            except Exception as e:
                print(f"Error merging Kundli PDF: {e}")

        return filepath

    def generate_docx(self, form_data):
        """Generate Word Destiny Report from form data"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name = self.sanitize_input(form_data.get('name', 'Client')).replace(' ', '_')
        filename = f'destiny_report_{client_name}_{timestamp}.docx'
        filepath = os.path.join(self.output_dir, filename)

        # Create document
        doc = Document()

        # Add main title with Times New Roman
        title = doc.add_heading('DESTINY REPORT', 0)
        title.alignment = 1  # Center
        for run in title.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(28)
        doc.add_paragraph()

        # Helper function to add orange border line
        def add_orange_line():
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')  # 12/8 = 1.5pt line thickness
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'FF8C00')  # Orange color
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Helper function to add orange box borders around paragraphs
        def add_orange_box_border(paragraph):
            """Add orange box border around a paragraph"""
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            # Add borders on all sides
            for border_side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '16')  # Border thickness
                border.set(qn('w:space'), '4')
                border.set(qn('w:color'), 'FF8C00')  # Orange color
                pBdr.append(border)

            pPr.append(pBdr)

            # Add padding
            pPr = paragraph._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '120')
            spacing.set(qn('w:after'), '120')
            pPr.append(spacing)

        # Helper function to format multi-line values with bullet points
        def format_value_with_bullets_docx(value):
            """Format values that contain multiple lines or comma-separated items with bullets"""
            if not value:
                return []

            # Check if value contains newlines or commas suggesting multiple items
            if '\n' in value:
                items = [item.strip() for item in value.split('\n') if item.strip()]
            elif ',' in value and len(value) > 50:  # Only split by comma if text is long
                items = [item.strip() for item in value.split(',') if item.strip()]
            else:
                # Single value, no bullets needed
                return [value]

            return items if len(items) > 1 else [value]

        # Helper function to check if text already has bullet points
        def has_bullet_docx(text):
            """Check if text starts with a bullet point"""
            if not text:
                return False
            text = text.strip()
            # Check for common bullet characters
            bullet_chars = ['•', '◦', '○', '●', '▪', '▫', '–', '-', '*', '→']
            return any(text.startswith(char) for char in bullet_chars) or (len(text) > 2 and text[0].isdigit() and text[1] in ['.', ')'])

        # Helper function to add section (only if it has content)
        def add_section(title_text, fields):
            # Check if any field has a value
            has_content = any(self.sanitize_input(form_data.get(field_key, '')) for field_key, _ in fields)

            if not has_content:
                return  # Skip empty sections

            # Create a table to wrap the section with orange border
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False

            # Set table style with orange borders
            tbl = table._element
            tblPr = tbl._element.get_or_add_tblPr()
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '16')
                border.set(qn('w:color'), 'FF8C00')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            # Get the cell
            cell = table.rows[0].cells[0]

            # Set cell margins (padding)
            tcPr = cell._element.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '120')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)

            # Add section heading
            heading_p = cell.add_paragraph()
            heading_run = heading_p.add_run(title_text)
            heading_run.font.name = 'Times New Roman'
            heading_run.font.size = Pt(18)
            heading_run.font.color.rgb = RGBColor(255, 140, 0)
            heading_run.font.bold = True

            cell.add_paragraph()  # Add spacing

            # Add fields
            for field_key, field_label in fields:
                value = self.sanitize_input(form_data.get(field_key, ''))
                if value:
                    # Add field label
                    label_p = cell.add_paragraph()
                    label_run = label_p.add_run(f'{field_label}:')
                    label_run.bold = True
                    label_run.font.name = 'Times New Roman'
                    label_run.font.size = Pt(11)
                    label_run.font.color.rgb = RGBColor(51, 51, 51)

                    # Format value with bullets if multi-line
                    items = format_value_with_bullets_docx(value)
                    if len(items) > 1:
                        # Multiple items - use bullets only if not already present
                        for item in items:
                            if has_bullet_docx(item):
                                # Already has bullet, add as plain paragraph
                                p = cell.add_paragraph()
                                value_run = p.add_run(item)
                                value_run.font.name = 'Times New Roman'
                                value_run.font.size = Pt(11)
                            else:
                                # Add bullet manually (style doesn't work well in table cells)
                                p = cell.add_paragraph()
                                bullet_run = p.add_run('• ' + item)
                                bullet_run.font.name = 'Times New Roman'
                                bullet_run.font.size = Pt(11)
                    else:
                        # Single item - no bullet
                        p = cell.add_paragraph()
                        value_run = p.add_run(items[0])
                        value_run.font.name = 'Times New Roman'
                        value_run.font.size = Pt(11)

                    cell.add_paragraph()  # Add spacing after each field

            doc.add_paragraph()  # Minimal spacing after section

        # ABOUT THE CLIENT
        add_section("ABOUT THE CLIENT", [
            ('name', 'Name'),
            ('dateOfBirth', 'Date of Birth'),
            ('timeOfBirth', 'Time of Birth'),
            ('placeOfBirth', 'Place of Birth')
        ])

        # VASTU ANALYSIS
        add_section("VASTU ANALYSIS", [
            ('mapOfHouse', 'Map of the House'),
            ('vastuAnalysis', 'Analysis (Entrances, Kitchen, Washrooms)'),
            ('vastuRemedies', 'Remedies')
        ])

        # ASTROLOGY (Custom handling for Dashas)
        # Check if ASTROLOGY section has any content
        mahadasha = self.format_dasha(form_data, 'mahadasha')
        antardasha = self.format_dasha(form_data, 'antardasha')
        pratyantardasha = self.format_dasha(form_data, 'pratyantardasha')

        astrology_fields = [
            'donationsToDo', 'donationsToWhom', 'gemstones', 'mantra',
            'birthNakshatra', 'mobileDisplayPicture', 'beneficialSymbols',
            'nakshatraProsperitySymbols', 'nakshatraMentalPhysicalWellbeing',
            'nakshatraAccomplishments', 'nakshatraAvoidSymbols'
        ]

        has_astrology_content = (
            mahadasha or antardasha or pratyantardasha or
            any(self.sanitize_input(form_data.get(field, '')) for field in astrology_fields)
        )

        if has_astrology_content:
            # Create a table to wrap the section with orange border
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False

            # Set table style with orange borders
            tbl = table._element
            tblPr = tbl._element.get_or_add_tblPr()
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '16')
                border.set(qn('w:color'), 'FF8C00')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            # Get the cell
            cell = table.rows[0].cells[0]

            # Set cell margins (padding)
            tcPr = cell._element.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '120')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)

            # Add section heading
            heading_p = cell.add_paragraph()
            heading_run = heading_p.add_run("ASTROLOGY")
            heading_run.font.name = 'Times New Roman'
            heading_run.font.size = Pt(18)
            heading_run.font.color.rgb = RGBColor(255, 140, 0)
            heading_run.font.bold = True

            cell.add_paragraph()  # Add spacing

            # Add Mahadasha
            if mahadasha:
                label_p = cell.add_paragraph()
                label_run = label_p.add_run('Mahadasha:')
                label_run.bold = True
                label_run.font.name = 'Times New Roman'
                label_run.font.size = Pt(11)
                label_run.font.color.rgb = RGBColor(51, 51, 51)

                p = cell.add_paragraph()
                value_run = p.add_run(mahadasha)
                value_run.font.name = 'Times New Roman'
                value_run.font.size = Pt(11)
                cell.add_paragraph()

            # Add Antardasha
            if antardasha:
                label_p = cell.add_paragraph()
                label_run = label_p.add_run('Antardasha:')
                label_run.bold = True
                label_run.font.name = 'Times New Roman'
                label_run.font.size = Pt(11)
                label_run.font.color.rgb = RGBColor(51, 51, 51)

                p = cell.add_paragraph()
                value_run = p.add_run(antardasha)
                value_run.font.name = 'Times New Roman'
                value_run.font.size = Pt(11)
                cell.add_paragraph()

            # Add Pratyantardasha
            if pratyantardasha:
                label_p = cell.add_paragraph()
                label_run = label_p.add_run('Pratyantar Dasha:')
                label_run.bold = True
                label_run.font.name = 'Times New Roman'
                label_run.font.size = Pt(11)
                label_run.font.color.rgb = RGBColor(51, 51, 51)

                p = cell.add_paragraph()
                value_run = p.add_run(pratyantardasha)
                value_run.font.name = 'Times New Roman'
                value_run.font.size = Pt(11)
                cell.add_paragraph()

            # Add other astrology fields
            for field_key, field_label in [
                ('donationsToDo', 'Donations to Do'),
                ('donationsToWhom', 'Donations - To Whom'),
                ('gemstones', 'Gemstones'),
                ('mantra', 'Mantra to Make Situation Positive'),
                ('birthNakshatra', 'Your Birth Nakshatra'),
                ('mobileDisplayPicture', 'Beneficial Mobile Display Picture'),
                ('beneficialSymbols', 'Most Beneficial Symbols'),
                ('nakshatraProsperitySymbols', 'Prosperity Giving Symbols'),
                ('nakshatraMentalPhysicalWellbeing', 'Mental/Physical Wellbeing Symbols'),
                ('nakshatraAccomplishments', 'Accomplishment/Achievement Symbols'),
                ('nakshatraAvoidSymbols', 'Symbols to Avoid')
            ]:
                value = self.sanitize_input(form_data.get(field_key, ''))
                if value:
                    # Add field label
                    label_p = cell.add_paragraph()
                    label_run = label_p.add_run(f'{field_label}:')
                    label_run.bold = True
                    label_run.font.name = 'Times New Roman'
                    label_run.font.size = Pt(11)
                    label_run.font.color.rgb = RGBColor(51, 51, 51)

                    # Format value with bullets if multi-line
                    items = format_value_with_bullets_docx(value)
                    if len(items) > 1:
                        # Multiple items - use bullets only if not already present
                        for item in items:
                            if has_bullet_docx(item):
                                # Already has bullet, add as plain paragraph
                                p = cell.add_paragraph()
                                value_run = p.add_run(item)
                                value_run.font.name = 'Times New Roman'
                                value_run.font.size = Pt(11)
                            else:
                                # Add bullet manually (style doesn't work well in table cells)
                                p = cell.add_paragraph()
                                bullet_run = p.add_run('• ' + item)
                                bullet_run.font.name = 'Times New Roman'
                                bullet_run.font.size = Pt(11)
                    else:
                        # Single item - no bullet
                        p = cell.add_paragraph()
                        value_run = p.add_run(items[0])
                        value_run.font.name = 'Times New Roman'
                        value_run.font.size = Pt(11)

                    cell.add_paragraph()  # Add spacing after each field

            doc.add_paragraph()  # Minimal spacing after section

        # ASTRO VASTU SOLUTION
        add_section("ASTRO VASTU SOLUTION", [
            ('aspectsOnHouses', 'Aspects on Houses'),
            ('aspectsOnPlanets', 'Aspects on Planets'),
            ('whatToRemove', 'What to Remove from Which Directions'),
            ('whatToPlace', 'What to Place in Which Directions'),
            ('astroVastuRemediesBody', 'Astro Vastu Remedies for Body'),
            ('colorObjectsToUse', 'What Color or Objects to Use on Body'),
            ('colorObjectsNotToUse', 'What Color or Objects NOT to Use on Body'),
            ('lockerLocation', 'Most Favorable Location of Locker'),
            ('laughingBuddhaDirection', 'Placement of Laughing Buddha/Vision Board - Direction'),
            ('wishListInkColor', 'Color of Ink to Write Wish List')
        ])

        # GUIDELINES
        add_section("GUIDELINES", [
            ('neverCriticize', 'To Whom You Should Never Criticize or Judge'),
            ('importantBooks', 'Important Books to Read to Uplift Your Life'),
            ('giftsToGive', 'Gifts - To Give'),
            ('giftsToReceive', 'Gifts - To Receive')
        ])

        # BHRIGUNANDA NADI
        add_section("BHRIGUNANDA NADI", [
            ('saturnRelation', 'Saturn Relation'),
            ('saturnFollowing', 'Saturn is following'),
            ('professionalMindset', 'Professional Mindset'),
            ('venusRelation', 'Venus Relation'),
            ('venusFollowing', 'Venus is following'),
            ('financialMindset', 'Financial Mindset')
        ])

        # Add timestamp
        timestamp_para = doc.add_paragraph(
            f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        timestamp_para.runs[0].italic = True

        # Save document
        doc.save(filepath)
        return filepath

    def generate_excel(self, form_data):
        """Generate Excel Destiny Report from form data"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name = self.sanitize_input(form_data.get('name', 'Client')).replace(' ', '_')
        filename = f'destiny_report_{client_name}_{timestamp}.xlsx'
        filepath = os.path.join(self.output_dir, filename)

        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Destiny Report"

        # Add main title
        ws['A1'] = 'DESTINY REPORT'
        ws['A1'].font = ws['A1'].font.copy(size=18, bold=True)
        ws.merge_cells('A1:B1')

        row = 3

        # Helper function to add section
        def add_section(title, fields):
            nonlocal row
            # Section title
            ws[f'A{row}'] = title
            ws[f'A{row}'].font = ws[f'A{row}'].font.copy(size=14, bold=True)
            ws.merge_cells(f'A{row}:B{row}')
            row += 1

            # Section fields
            for field_key, field_label in fields:
                value = self.sanitize_input(form_data.get(field_key, ''))
                if value:
                    ws[f'A{row}'] = field_label
                    ws[f'B{row}'] = str(value)
                    ws[f'A{row}'].font = ws[f'A{row}'].font.copy(bold=True)
                    row += 1

            row += 1  # Add space after section

        # ABOUT THE CLIENT
        add_section("ABOUT THE CLIENT", [
            ('name', 'Name'),
            ('dateOfBirth', 'Date of Birth'),
            ('timeOfBirth', 'Time of Birth'),
            ('placeOfBirth', 'Place of Birth')
        ])

        # VASTU ANALYSIS
        add_section("VASTU ANALYSIS", [
            ('mapOfHouse', 'Map of the House'),
            ('vastuAnalysis', 'Analysis (Entrances, Kitchen, Washrooms)'),
            ('vastuRemedies', 'Remedies')
        ])

        # ASTROLOGY
        ws[f'A{row}'] = "ASTROLOGY"
        ws[f'A{row}'].font = ws[f'A{row}'].font.copy(size=14, bold=True)
        ws.merge_cells(f'A{row}:B{row}')
        row += 1

        # Add Mahadasha
        mahadasha = self.format_dasha(form_data, 'mahadasha')
        if mahadasha:
            ws[f'A{row}'] = 'Mahadasha'
            ws[f'B{row}'] = mahadasha
            ws[f'A{row}'].font = ws[f'A{row}'].font.copy(bold=True)
            row += 1

        # Add Antardasha
        antardasha = self.format_dasha(form_data, 'antardasha')
        if antardasha:
            ws[f'A{row}'] = 'Antardasha'
            ws[f'B{row}'] = antardasha
            ws[f'A{row}'].font = ws[f'A{row}'].font.copy(bold=True)
            row += 1

        # Add Pratyantardasha
        pratyantardasha = self.format_dasha(form_data, 'pratyantardasha')
        if pratyantardasha:
            ws[f'A{row}'] = 'Pratyantar Dasha'
            ws[f'B{row}'] = pratyantardasha
            ws[f'A{row}'].font = ws[f'A{row}'].font.copy(bold=True)
            row += 1

        # Add other astrology fields
        for field_key, field_label in [
            ('donationsToDo', 'Donations to Do'),
            ('donationsToWhom', 'Donations - To Whom'),
            ('gemstones', 'Gemstones'),
            ('mantra', 'Mantra to Make Situation Positive'),
            ('birthNakshatra', 'Your Birth Nakshatra'),
            ('mobileDisplayPicture', 'Beneficial Mobile Display Picture'),
            ('beneficialSymbols', 'Most Beneficial Symbols'),
            ('nakshatraProsperitySymbols', 'Prosperity Giving Symbols'),
            ('nakshatraMentalPhysicalWellbeing', 'Mental/Physical Wellbeing Symbols'),
            ('nakshatraAccomplishments', 'Accomplishment/Achievement Symbols'),
            ('nakshatraAvoidSymbols', 'Symbols to Avoid')
        ]:
            value = self.sanitize_input(form_data.get(field_key, ''))
            if value:
                ws[f'A{row}'] = field_label
                ws[f'B{row}'] = str(value)
                ws[f'A{row}'].font = ws[f'A{row}'].font.copy(bold=True)
                row += 1

        row += 1

        # ASTRO VASTU SOLUTION
        add_section("ASTRO VASTU SOLUTION", [
            ('aspectsOnHouses', 'Aspects on Houses'),
            ('aspectsOnPlanets', 'Aspects on Planets'),
            ('whatToRemove', 'What to Remove from Which Directions'),
            ('whatToPlace', 'What to Place in Which Directions'),
            ('astroVastuRemediesBody', 'Astro Vastu Remedies for Body'),
            ('colorObjectsToUse', 'What Color or Objects to Use on Body'),
            ('colorObjectsNotToUse', 'What Color or Objects NOT to Use on Body'),
            ('lockerLocation', 'Most Favorable Location of Locker'),
            ('laughingBuddhaDirection', 'Placement of Laughing Buddha/Vision Board - Direction'),
            ('wishListInkColor', 'Color of Ink to Write Wish List')
        ])

        # GUIDELINES
        add_section("GUIDELINES", [
            ('neverCriticize', 'To Whom You Should Never Criticize or Judge'),
            ('importantBooks', 'Important Books to Read to Uplift Your Life'),
            ('giftsToGive', 'Gifts - To Give'),
            ('giftsToReceive', 'Gifts - To Receive')
        ])

        # BHRIGUNANDA NADI
        add_section("BHRIGUNANDA NADI", [
            ('saturnRelation', 'Saturn Relation'),
            ('saturnFollowing', 'Saturn is following'),
            ('professionalMindset', 'Professional Mindset'),
            ('venusRelation', 'Venus Relation'),
            ('venusFollowing', 'Venus is following'),
            ('financialMindset', 'Financial Mindset')
        ])

        # Add timestamp
        ws[f'A{row}'] = f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        ws[f'A{row}'].font = ws[f'A{row}'].font.copy(italic=True)

        # Adjust column widths
        ws.column_dimensions['A'].width = 40
        ws.column_dimensions['B'].width = 60

        # Save workbook
        wb.save(filepath)
        return filepath
